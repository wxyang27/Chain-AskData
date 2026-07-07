import re
from pathlib import Path
from typing import Any

from docx import Document

from app.knowledge_importer.excel_loader import clean_text


TABLE_NAME_PATTERN = re.compile(r"\b(?:dim|dwd|dws|dm)_[a-zA-Z0-9_]+\b")


def infer_layer(table_name: str) -> str:
    if table_name.startswith("dim_"):
        return "DIM"
    if table_name.startswith("dwd_"):
        return "DWD"
    if table_name.startswith("dws_"):
        return "DWS"
    if table_name.startswith("dm_"):
        return "DM"
    return "UNKNOWN"


def table_to_text(table: Any) -> str:
    lines = []
    for row in table.rows:
        cells = [clean_text(cell.text) for cell in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def iter_document_text(document: Any) -> list[str]:
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            parts.append(text)
    for table in document.tables:
        text = table_to_text(table)
        if text:
            parts.append(text)
    return parts


class DatabaseDocxLoader:
    """Extract table assets from the uploaded database Word document."""

    def load_tables(
        self, docx_path: Path, data_sources: list[dict[str, Any]]
    ) -> list[dict[str, Any]]:
        document = Document(docx_path)
        all_text = "\n".join(iter_document_text(document))
        table_names = sorted(set(TABLE_NAME_PATTERN.findall(all_text)))
        data_source_by_table = {source["table_name"]: source for source in data_sources}

        tables = []
        for index, table_name in enumerate(table_names, start=1):
            data_source = data_source_by_table.get(table_name, {})
            tables.append(
                {
                    "asset_id": f"table:{table_name}",
                    "asset_type": "table",
                    "table_name": table_name,
                    "full_name": f"soyoung_dw.{table_name}",
                    "layer": infer_layer(table_name),
                    "theme": data_source.get("business_description", ""),
                    "grain": data_source.get("notes", ""),
                    "partition_field": data_source.get("partition_field", "dp"),
                    "main_fields": data_source.get("main_fields", []),
                    "business_description": data_source.get("business_description", ""),
                    "source_file": docx_path.name,
                    "source_index": index,
                }
            )

        return tables


class AnalysisDocxLoader:
    """Extract business playbook chunks from the uploaded analysis Word document."""

    def load_playbooks(self, docx_path: Path) -> list[dict[str, Any]]:
        document = Document(docx_path)
        playbooks: list[dict[str, Any]] = []

        playbooks.extend(self._paragraph_sections(docx_path, document))
        playbooks.extend(self._table_sections(docx_path, document))
        return playbooks

    def _paragraph_sections(self, docx_path: Path, document: Any) -> list[dict[str, Any]]:
        sections: list[dict[str, Any]] = []
        current_title = "连锁经营分析方法"
        current_lines: list[str] = []

        for paragraph in document.paragraphs:
            text = clean_text(paragraph.text)
            if not text:
                continue
            if self._looks_like_heading(text) and current_lines:
                sections.append(
                    self._make_playbook(
                        docx_path=docx_path,
                        index=len(sections) + 1,
                        title=current_title,
                        content="\n".join(current_lines),
                        content_type="section",
                    )
                )
                current_title = text[:80]
                current_lines = [text]
            else:
                if self._looks_like_heading(text):
                    current_title = text[:80]
                current_lines.append(text)

        if current_lines:
            sections.append(
                self._make_playbook(
                    docx_path=docx_path,
                    index=len(sections) + 1,
                    title=current_title,
                    content="\n".join(current_lines),
                    content_type="section",
                )
            )

        return [section for section in sections if len(section["content"]) >= 20]

    def _table_sections(self, docx_path: Path, document: Any) -> list[dict[str, Any]]:
        sections: list[dict[str, Any]] = []
        for index, table in enumerate(document.tables, start=1):
            content = table_to_text(table)
            if not content:
                continue
            title = content.splitlines()[0][:80]
            sections.append(
                self._make_playbook(
                    docx_path=docx_path,
                    index=index,
                    title=title,
                    content=content,
                    content_type="table",
                )
            )
        return sections

    def _make_playbook(
        self,
        docx_path: Path,
        index: int,
        title: str,
        content: str,
        content_type: str,
    ) -> dict[str, Any]:
        return {
            "asset_id": f"business_playbook:{content_type}:{index:03d}",
            "asset_type": "business_playbook",
            "playbook_id": f"{content_type}_{index:03d}",
            "title": title,
            "content_type": content_type,
            "content": content,
            "source_file": docx_path.name,
            "source_index": index,
        }

    def _looks_like_heading(self, text: str) -> bool:
        if len(text) > 80:
            return False
        return bool(
            re.match(r"^(第[一二三四五六七八九十]+[章节]|[0-9]+[.、]|[一二三四五六七八九十]+、)", text)
        )

